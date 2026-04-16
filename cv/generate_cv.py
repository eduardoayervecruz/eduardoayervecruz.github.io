#!/usr/bin/env python3
"""
CV Generator - Reads cv_data.json and generates CVs in English, Spanish, and Korean.
Outputs: Word (.docx), PDF, and HTML for each language (9 files total).

All translations are manually verified in cv_data.json (no machine translation).

Usage:
    python generate_cv.py

Requirements:
    pip install python-docx reportlab
"""

import json
import os
import shutil
import sys
from pathlib import Path

# --- Document generation ---
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether, Image
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ============================================================
# CONFIGURATION
# ============================================================

SCRIPT_DIR = Path(__file__).parent
DATA_FILE = SCRIPT_DIR / "cv_data.json"
OUTPUT_DIR = SCRIPT_DIR / "output"
PORTRAIT_FILE = SCRIPT_DIR / "portrait.JPG"

LANGUAGES = {
    "en": "English",
    "es": "Spanish",
    "ko": "Korean"
}

# Section header translations (pre-defined for consistency)
SECTION_HEADERS = {
    "en": {
        "summary": "Professional Summary",
        "experience": "Experience",
        "projects": "Featured Projects",
        "tech_stack": "Tech Stack",
        "education": "Education",
        "certifications": "Certifications",
        "languages": "Languages",
        "awards": "Awards & Honors",
        "extracurricular": "Extracurricular Activities"
    },
    "es": {
        "summary": "Resumen Profesional",
        "experience": "Experiencia",
        "projects": "Proyectos Destacados",
        "tech_stack": "Stack Tecnológico",
        "education": "Educación",
        "certifications": "Certificaciones",
        "languages": "Idiomas",
        "awards": "Premios y Reconocimientos",
        "extracurricular": "Actividades Extracurriculares"
    },
    "ko": {
        "summary": "전문 요약",
        "experience": "경력",
        "projects": "주요 프로젝트",
        "tech_stack": "기술 스택",
        "education": "학력",
        "certifications": "자격증 및 시험",
        "languages": "언어",
        "awards": "수상 및 장학금",
        "extracurricular": "대외 활동"
    }
}

TECH_STACK_LABELS = {
    "en": {
        "languages": "Languages",
        "frontend": "Frontend",
        "backend": "Backend",
        "databases": "Databases",
        "devops": "DevOps & Cloud",
        "tools": "Tools"
    },
    "es": {
        "languages": "Lenguajes",
        "frontend": "Frontend",
        "backend": "Backend",
        "databases": "Bases de Datos",
        "devops": "DevOps y Cloud",
        "tools": "Herramientas"
    },
    "ko": {
        "languages": "프로그래밍 언어",
        "frontend": "프론트엔드",
        "backend": "백엔드",
        "databases": "데이터베이스",
        "devops": "데브옵스 & 클라우드",
        "tools": "도구"
    }
}

# Color scheme
PRIMARY_COLOR = "#2B547E"    # Dark blue
ACCENT_COLOR = "#4A90D9"     # Medium blue
TEXT_COLOR = "#333333"        # Dark gray
LIGHT_GRAY = "#F5F5F5"

# ============================================================
# TRANSLATION FUNCTIONS (manual overrides only, no API)
# ============================================================

def _get(obj, field, lang):
    """Get a translated field from an object. Falls back to the English field."""
    if lang == "en":
        return obj.get(field, "")
    translated_key = f"{field}_{lang}"
    val = obj.get(translated_key)
    if val is not None:
        return val
    # Fallback to English
    return obj.get(field, "")


def translate_cv_data(data, lang):
    """Build a translated CV data structure using manual _es/_ko fields."""
    if lang == "en":
        return data

    print(f"  Applying {LANGUAGES[lang]} translations...")
    translated = json.loads(json.dumps(data))  # Deep copy

    # Personal info
    translated["personal"]["title"] = _get(data["personal"], "title", lang)
    translated["personal"]["location"] = _get(data["personal"], "location", lang)

    # Summary
    translated["summary"] = _get(data, "summary", lang)

    # Experience
    for i, exp in enumerate(data["experience"]):
        translated["experience"][i]["role"] = _get(exp, "role", lang)
        translated["experience"][i]["end_date"] = _get(exp, "end_date", lang)

        override_key = f"highlights_{lang}"
        overrides = exp.get(override_key, [])
        translated_highlights = []
        for j, h in enumerate(exp["highlights"]):
            manual = overrides[j] if j < len(overrides) else None
            if manual:
                translated_highlights.append(manual)
            else:
                print(f"    ⚠ Missing {lang} translation: experience[{i}].highlights[{j}]")
                translated_highlights.append(h)  # fallback to English
        translated["experience"][i]["highlights"] = translated_highlights

    # Featured projects
    for i, proj in enumerate(data["featured_projects"]):
        translated["featured_projects"][i]["name"] = _get(proj, "name", lang)
        translated["featured_projects"][i]["description"] = _get(proj, "description", lang)

    # Education
    for i, edu in enumerate(data["education"]):
        translated["education"][i]["degree"] = _get(edu, "degree", lang)
        if edu.get("notes"):
            translated["education"][i]["notes"] = _get(edu, "notes", lang)

    # Certifications - keep as-is (proper nouns / test names)

    # Awards
    awards_key = f"awards_{lang}"
    if data.get(awards_key):
        translated["awards"] = data[awards_key]
    elif data.get("awards"):
        print(f"    ⚠ Missing {lang} translations for awards")

    # Extracurricular
    if data.get("extracurricular"):
        for i, ext in enumerate(data["extracurricular"]):
            translated["extracurricular"][i]["organization"] = _get(ext, "organization", lang)
            translated["extracurricular"][i]["role"] = _get(ext, "role", lang)
            translated["extracurricular"][i]["period"] = _get(ext, "period", lang)

            override_key = f"highlights_{lang}"
            overrides = ext.get(override_key, [])
            translated_highlights = []
            for j, h in enumerate(ext.get("highlights", [])):
                manual = overrides[j] if j < len(overrides) else None
                if manual:
                    translated_highlights.append(manual)
                else:
                    print(f"    ⚠ Missing {lang} translation: extracurricular[{i}].highlights[{j}]")
                    translated_highlights.append(h)
            translated["extracurricular"][i]["highlights"] = translated_highlights

    # Languages spoken
    if data.get("languages_spoken"):
        for i, spoken in enumerate(data["languages_spoken"]):
            translated["languages_spoken"][i]["language"] = _get(spoken, "language", lang)
            translated["languages_spoken"][i]["level"] = _get(spoken, "level", lang)

    return translated


# ============================================================
# WORD (.docx) GENERATOR
# ============================================================

def generate_docx(data, lang, output_path):
    """Generate a Word document CV."""
    doc = Document()
    headers = SECTION_HEADERS[lang]
    ts_labels = TECH_STACK_LABELS[lang]

    # --- Page margins (balanced for 2-page fit) ---
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(12)

    # --- Header with portrait (table: left=info, right=photo) ---
    p = data["personal"]
    from docx.oxml.ns import qn

    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    # Set column widths: left ~80%, right ~20% (18cm usable with 1.5cm margins)
    header_table.columns[0].width = Cm(14)
    header_table.columns[1].width = Cm(4)

    # Remove table borders
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(b)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Left cell: name, title, contact
    left_cell = header_table.cell(0, 0)
    # Remove default empty paragraph
    left_cell.paragraphs[0].clear()

    name_para = left_cell.paragraphs[0]
    name_run = name_para.add_run(data["personal"]["name"])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    name_para.space_after = Pt(1)

    title_para = left_cell.add_paragraph()
    title_run = title_para.add_run(data["personal"]["title"])
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    title_para.space_after = Pt(2)

    contact_fields = ["email", "phone", "location", "linkedin", "github", "portfolio"]
    for key in contact_fields:
        if p.get(key):
            contact_para = left_cell.add_paragraph()
            contact_run = contact_para.add_run(p[key])
            contact_run.font.size = Pt(8)
            contact_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            contact_para.space_after = Pt(0)
            contact_para.space_before = Pt(0)

    # Right cell: portrait
    right_cell = header_table.cell(0, 1)
    right_cell.paragraphs[0].clear()
    right_cell.vertical_alignment = 0  # TOP
    img_para = right_cell.paragraphs[0]
    img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if PORTRAIT_FILE.exists():
        img_run = img_para.add_run()
        img_run.add_picture(str(PORTRAIT_FILE), width=Cm(3.0))

    # Add spacing after header
    spacer = doc.add_paragraph()
    spacer.space_before = Pt(4)
    spacer.space_after = Pt(0)

    def add_section_heading(text):
        heading = doc.add_paragraph()
        run = heading.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading.space_before = Pt(8)
        heading.space_after = Pt(3)
        from docx.oxml.ns import qn
        pBdr = heading._p.get_or_add_pPr()
        bottom = pBdr.makeelement(qn('w:pBdr'), {})
        b = bottom.makeelement(qn('w:bottom'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '1',
            qn('w:color'): '000000'
        })
        bottom.append(b)
        pBdr.append(bottom)

    # --- Summary ---
    add_section_heading(headers["summary"])
    summary_para = doc.add_paragraph(data["summary"])
    summary_para.paragraph_format.space_after = Pt(4)

    # --- Experience ---
    add_section_heading(headers["experience"])
    for exp in data["experience"]:
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(f"{exp['role']}  —  {exp['company']}")
        role_run.bold = True
        role_run.font.size = Pt(9)
        role_para.space_after = Pt(0)

        date_para = doc.add_paragraph()
        date_run = date_para.add_run(
            f"{exp['start_date']} – {exp['end_date']}  |  {exp['location']}"
        )
        date_run.font.size = Pt(8)
        date_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        date_para.space_after = Pt(2)

        for highlight in exp["highlights"]:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.text = highlight
            bullet.paragraph_format.space_after = Pt(1)

    # --- Featured Projects ---
    add_section_heading(headers["projects"])
    for proj in data["featured_projects"]:
        proj_para = doc.add_paragraph()
        name_run = proj_para.add_run(proj["name"])
        name_run.bold = True
        name_run.font.size = Pt(9)
        if proj.get("link"):
            link_run = proj_para.add_run(f"  ({proj['link']})")
            link_run.font.size = Pt(8)
        proj_para.space_after = Pt(0)

        desc_para = doc.add_paragraph(proj["description"])
        desc_para.space_after = Pt(0)

        tech_para = doc.add_paragraph()
        tech_label = tech_para.add_run("Tech: ")
        tech_label.bold = True
        tech_label.font.size = Pt(8)
        tech_para.add_run(proj["tech_stack"]).font.size = Pt(8)
        tech_para.space_after = Pt(4)

    # --- Tech Stack ---
    add_section_heading(headers["tech_stack"])
    ts = data["tech_stack"]
    for key, label in ts_labels.items():
        if ts.get(key):
            ts_para = doc.add_paragraph()
            label_run = ts_para.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(8.5)
            ts_run = ts_para.add_run(ts[key])
            ts_run.font.size = Pt(8.5)
            ts_para.paragraph_format.space_after = Pt(0)

    # --- Education ---
    add_section_heading(headers["education"])
    for edu in data["education"]:
        edu_para = doc.add_paragraph()
        deg_run = edu_para.add_run(edu["degree"])
        deg_run.bold = True
        deg_run.font.size = Pt(9)
        inst_run = edu_para.add_run(f"  —  {edu['institution']}")
        inst_run.font.size = Pt(9)
        edu_para.space_after = Pt(0)

        details = doc.add_paragraph()
        detail_run = details.add_run(f"{edu['year']}  |  {edu['location']}")
        detail_run.font.size = Pt(8)
        detail_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if edu.get("notes"):
            details.add_run(f"\n{edu['notes']}").font.size = Pt(8)
        details.space_after = Pt(3)

    # --- Awards ---
    if data.get("awards"):
        add_section_heading(headers["awards"])
        for award in data["awards"]:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.text = award
            bullet.paragraph_format.space_after = Pt(0)

    # --- Certifications ---
    if data.get("certifications"):
        add_section_heading(headers["certifications"])
        for cert in data["certifications"]:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.text = cert
            bullet.paragraph_format.space_after = Pt(0)

    # --- Extracurricular ---
    if data.get("extracurricular"):
        add_section_heading(headers["extracurricular"])
        for ext in data["extracurricular"]:
            ext_para = doc.add_paragraph()
            role_run = ext_para.add_run(f"{ext['role']}  —  {ext['organization']}")
            role_run.bold = True
            role_run.font.size = Pt(9)
            ext_para.space_after = Pt(0)

            period_para = doc.add_paragraph()
            period_run = period_para.add_run(ext['period'])
            period_run.font.size = Pt(8)
            period_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            period_para.space_after = Pt(1)

            for h in ext.get("highlights", []):
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.text = h
                bullet.paragraph_format.space_after = Pt(0)

    # --- Languages (each on its own line) ---
    if data.get("languages_spoken"):
        add_section_heading(headers["languages"])
        for spoken in data["languages_spoken"]:
            lang_para = doc.add_paragraph()
            label_run = lang_para.add_run(f"{spoken['language']}: ")
            label_run.bold = True
            label_run.font.size = Pt(8.5)
            lang_run = lang_para.add_run(spoken['level'])
            lang_run.font.size = Pt(8.5)
            lang_para.paragraph_format.space_after = Pt(0)

    doc.save(str(output_path))
    print(f"  ✓ Word: {output_path.name}")


# ============================================================
# PDF GENERATOR
# ============================================================

def _setup_pdf_fonts(lang):
    """Register fonts and return (regular, bold) font names."""
    if lang == 'ko':
        # 1) Try CID font (built into ReportLab, works on all platforms)
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        try:
            pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))
            return 'HYSMyeongJo-Medium', 'HYSMyeongJo-Medium'
        except Exception:
            pass

        # 2) Try platform-specific TTF fonts
        font_candidates = [
            # macOS
            ("/System/Library/Fonts/Supplemental/AppleGothic.ttf", None),
            # Linux - user local
            (os.path.expanduser("~/.local/share/fonts/NanumGothic-Regular.ttf"),
             os.path.expanduser("~/.local/share/fonts/NanumGothic-Bold.ttf")),
            # Linux - system
            ("/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
             "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf"),
        ]
        for reg, bold in font_candidates:
            if os.path.exists(reg):
                try:
                    pdfmetrics.registerFont(TTFont('KoreanFont', reg))
                    bold_path = bold if bold and os.path.exists(bold) else reg
                    pdfmetrics.registerFont(TTFont('KoreanFont-Bold', bold_path))
                    return 'KoreanFont', 'KoreanFont-Bold'
                except Exception:
                    continue

        print("  Warning: No Korean font found, Korean text may not render.")

    return 'Helvetica', 'Helvetica-Bold'


def generate_pdf(data, lang, output_path):
    """Generate a PDF CV using ReportLab."""
    base_font, bold_font = _setup_pdf_fonts(lang)
    headers = SECTION_HEADERS[lang]
    ts_labels = TECH_STACK_LABELS[lang]
    BLACK = HexColor('#000000')
    DARK = HexColor('#333333')

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.4*inch,
        bottomMargin=0.4*inch,
        leftMargin=0.6*inch,
        rightMargin=0.6*inch,
    )

    styles = getSampleStyleSheet()

    # ── Styles (balanced for full 2-page fit) ──
    s_name = ParagraphStyle(
        'CVName', parent=styles['Title'],
        fontName=bold_font, fontSize=20, textColor=BLACK,
        alignment=0, spaceAfter=1, spaceBefore=0, leading=24
    )
    s_title = ParagraphStyle(
        'CVTitle', parent=styles['Normal'],
        fontName=bold_font, fontSize=10.5, textColor=DARK,
        alignment=0, spaceAfter=4, leading=13
    )
    s_contact = ParagraphStyle(
        'CVContact', parent=styles['Normal'],
        fontName=base_font, fontSize=8.5, textColor=DARK,
        alignment=0, spaceAfter=0, leading=12
    )
    s_section = ParagraphStyle(
        'CVSection', parent=styles['Heading2'],
        fontName=bold_font, fontSize=10.5, textColor=BLACK,
        spaceBefore=8, spaceAfter=2, leading=13,
    )
    s_body = ParagraphStyle(
        'CVBody', parent=styles['Normal'],
        fontName=base_font, fontSize=9, textColor=DARK,
        spaceAfter=3, leading=12.5
    )
    s_role = ParagraphStyle(
        'CVRole', parent=styles['Normal'],
        fontName=bold_font, fontSize=9.5, textColor=BLACK,
        spaceAfter=1, leading=12
    )
    s_meta = ParagraphStyle(
        'CVMeta', parent=styles['Normal'],
        fontName=base_font, fontSize=8.5, textColor=DARK,
        spaceAfter=2, leading=11
    )
    s_bullet = ParagraphStyle(
        'CVBullet', parent=styles['Normal'],
        fontName=base_font, fontSize=9, textColor=DARK,
        leftIndent=14, spaceAfter=1.5, leading=12,
        bulletIndent=4, bulletFontSize=8.5
    )
    s_item = ParagraphStyle(
        'CVItem', parent=styles['Normal'],
        fontName=base_font, fontSize=9, textColor=DARK,
        leftIndent=10, spaceAfter=1.5, leading=12
    )

    story = []
    p = data["personal"]

    # ── Header with portrait ──
    # Build left side (name + title + contact)
    left_parts = []
    left_parts.append(Paragraph(p["name"], s_name))
    left_parts.append(Paragraph(p["title"], s_title))
    contact_fields = ["email", "phone", "location", "linkedin", "github", "portfolio"]
    for key in contact_fields:
        if p.get(key):
            left_parts.append(Paragraph(p[key], s_contact))

    # Build right side (portrait)
    if PORTRAIT_FILE.exists():
        portrait_img = Image(str(PORTRAIT_FILE), width=1.0*inch, height=1.25*inch)
        portrait_img.hAlign = 'RIGHT'
        header_table = Table(
            [[left_parts, portrait_img]],
            colWidths=[5.4*inch, 1.4*inch]
        )
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(header_table)
    else:
        story.extend(left_parts)
    story.append(Spacer(1, 3))

    def add_section_header(title):
        story.append(Paragraph(f"<b>{title.upper()}</b>", s_section))
        story.append(HRFlowable(
            width="100%", thickness=0.5, color=BLACK,
            spaceBefore=0, spaceAfter=4
        ))

    # ── Summary ──
    add_section_header(headers["summary"])
    story.append(Paragraph(data["summary"], s_body))

    # ── Experience ──
    add_section_header(headers["experience"])
    for i, exp in enumerate(data["experience"]):
        story.append(Paragraph(
            f"<b>{exp['role']}</b>  —  {exp['company']}", s_role
        ))
        story.append(Paragraph(
            f"{exp['start_date']} – {exp['end_date']}  |  {exp['location']}", s_meta
        ))
        for h in exp["highlights"]:
            story.append(Paragraph(f"• {h}", s_bullet))
        if i < len(data["experience"]) - 1:
            story.append(Spacer(1, 4))

    # ── Featured Projects ──
    add_section_header(headers["projects"])
    for i, proj in enumerate(data["featured_projects"]):
        link_text = f"  ({proj['link']})" if proj.get("link") else ""
        story.append(Paragraph(f"<b>{proj['name']}</b>{link_text}", s_role))
        story.append(Paragraph(proj["description"], s_body))
        story.append(Paragraph(f"<b>Tech:</b> {proj['tech_stack']}", s_meta))
        if i < len(data["featured_projects"]) - 1:
            story.append(Spacer(1, 3))

    # ── Tech Stack ──
    add_section_header(headers["tech_stack"])
    ts = data["tech_stack"]
    for key, label in ts_labels.items():
        if ts.get(key):
            story.append(Paragraph(f"<b>{label}:</b>  {ts[key]}", s_body))

    # ── Education ──
    add_section_header(headers["education"])
    for i, edu in enumerate(data["education"]):
        story.append(Paragraph(
            f"<b>{edu['degree']}</b>  —  {edu['institution']}", s_role
        ))
        meta_line = f"{edu['year']}  |  {edu['location']}"
        story.append(Paragraph(meta_line, s_meta))
        if edu.get("notes"):
            story.append(Paragraph(edu["notes"], s_item))
        if i < len(data["education"]) - 1:
            story.append(Spacer(1, 3))

    # ── Awards ──
    if data.get("awards"):
        add_section_header(headers["awards"])
        for award in data["awards"]:
            story.append(Paragraph(f"• {award}", s_bullet))

    # ── Certifications ──
    if data.get("certifications"):
        add_section_header(headers["certifications"])
        for cert in data["certifications"]:
            story.append(Paragraph(f"• {cert}", s_bullet))

    # ── Extracurricular ──
    if data.get("extracurricular"):
        add_section_header(headers["extracurricular"])
        for i, ext in enumerate(data["extracurricular"]):
            story.append(Paragraph(
                f"<b>{ext['role']}</b>  —  {ext['organization']}", s_role
            ))
            story.append(Paragraph(ext['period'], s_meta))
            for h in ext.get("highlights", []):
                story.append(Paragraph(f"• {h}", s_bullet))
            if i < len(data["extracurricular"]) - 1:
                story.append(Spacer(1, 2))

    # ── Languages ── (each on its own line)
    if data.get("languages_spoken"):
        add_section_header(headers["languages"])
        for spoken in data["languages_spoken"]:
            story.append(Paragraph(
                f"<b>{spoken['language']}:</b>  {spoken['level']}", s_body
            ))

    doc.build(story)
    print(f"  ✓ PDF:  {output_path.name}")


# ============================================================
# HTML GENERATOR
# ============================================================

def generate_html(data, lang, output_path):
    """Generate an HTML CV."""
    headers = SECTION_HEADERS[lang]
    ts_labels = TECH_STACK_LABELS[lang]
    p = data["personal"]

    html_lang = {"en": "en", "es": "es", "ko": "ko"}[lang]

    # Contact: each on its own line
    contact_items = []
    for key in ["email", "phone", "location", "linkedin", "github"]:
        if p.get(key):
            val = p[key]
            if key == "email":
                contact_items.append(f'<a href="mailto:{val}">{val}</a>')
            elif key in ("linkedin", "github"):
                url = val if val.startswith("http") else f"https://{val}"
                contact_items.append(f'<a href="{url}">{val}</a>')
            else:
                contact_items.append(val)
    contact_html = "<br>".join(contact_items)

    # Experience HTML
    exp_html = ""
    for exp in data["experience"]:
        highlights = "\n".join(f"<li>{h}</li>" for h in exp["highlights"])
        exp_html += f"""
        <div class="entry">
            <div class="entry-header">
                <span class="role">{exp['role']}  —  {exp['company']}</span>
            </div>
            <div class="entry-meta">{exp['start_date']} – {exp['end_date']}  |  {exp['location']}</div>
            <ul>{highlights}</ul>
        </div>"""

    # Projects HTML
    proj_html = ""
    for proj in data["featured_projects"]:
        link = ""
        if proj.get("link"):
            url = proj["link"] if proj["link"].startswith("http") else f"https://{proj['link']}"
            link = f' <a href="{url}" class="project-link">↗</a>'
        proj_html += f"""
        <div class="entry">
            <div class="entry-header">
                <span class="role">{proj['name']}{link}</span>
            </div>
            <p>{proj['description']}</p>
            <div class="tech-label"><strong>Tech:</strong> {proj['tech_stack']}</div>
        </div>"""

    # Tech stack HTML
    ts = data["tech_stack"]
    ts_html = ""
    for key, label in ts_labels.items():
        if ts.get(key):
            ts_html += f'<div class="stack-row"><strong>{label}:</strong> {ts[key]}</div>\n'

    # Education HTML
    edu_html = ""
    for edu in data["education"]:
        notes = f'<div class="entry-notes">{edu["notes"]}</div>' if edu.get("notes") else ""
        edu_html += f"""
        <div class="entry">
            <div class="entry-header">
                <span class="role">{edu['degree']}  —  {edu['institution']}</span>
            </div>
            <div class="entry-meta">{edu['year']}  |  {edu['location']}</div>
            {notes}
        </div>"""

    # Awards HTML
    awards_html = ""
    if data.get("awards"):
        award_items = "\n".join(f"<li>{a}</li>" for a in data["awards"])
        awards_html = f"""
    <section>
        <h2>{headers['awards']}</h2>
        <ul>{award_items}</ul>
    </section>"""

    # Certifications HTML
    cert_html = ""
    if data.get("certifications"):
        cert_items = "\n".join(f"<li>{c}</li>" for c in data["certifications"])
        cert_html = f"""
    <section>
        <h2>{headers['certifications']}</h2>
        <ul>{cert_items}</ul>
    </section>"""

    # Extracurricular HTML
    extra_html = ""
    if data.get("extracurricular"):
        extra_entries = ""
        for ext in data["extracurricular"]:
            highlights = "\n".join(f"<li>{h}</li>" for h in ext.get("highlights", []))
            extra_entries += f"""
        <div class="entry">
            <div class="entry-header">
                <span class="role">{ext['role']}  —  {ext['organization']}</span>
            </div>
            <div class="entry-meta">{ext['period']}</div>
            <ul>{highlights}</ul>
        </div>"""
        extra_html = f"""
    <section>
        <h2>{headers['extracurricular']}</h2>
        {extra_entries}
    </section>"""

    # Languages HTML (each on its own line)
    lang_html = ""
    if data.get("languages_spoken"):
        lang_items = "\n".join(
            f'<div class="lang-row"><strong>{l["language"]}:</strong> {l["level"]}</div>'
            for l in data["languages_spoken"]
        )
        lang_html = f"""
    <section>
        <h2>{headers['languages']}</h2>
        {lang_items}
    </section>"""

    html = f"""<!DOCTYPE html>
<html lang="{html_lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{p['name']} — CV</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
        @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700&display=swap');

        * {{ margin: 0; padding: 0; box-sizing: border-box; }}

        body {{
            font-family: 'Inter', 'Noto Sans KR', -apple-system, BlinkMacSystemFont, sans-serif;
            color: #333;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 30px;
            background: #fff;
        }}

        /* Header */
        .header {{ display: flex; align-items: flex-start; margin-bottom: 30px; gap: 20px; }}
        .header-info {{ flex: 1; }}
        .header-info h1 {{
            font-size: 28px; font-weight: 700;
            color: #000; margin-bottom: 4px;
        }}
        .header-info .title {{
            font-size: 14px; color: #333;
            font-weight: 600; margin-bottom: 8px;
        }}
        .header-info .contact {{
            font-size: 12px; color: #333; line-height: 1.8;
        }}
        .header-info .contact a {{
            color: #333; text-decoration: none;
        }}
        .header-info .contact a:hover {{ text-decoration: underline; }}
        .header-portrait {{
            flex-shrink: 0;
        }}
        .header-portrait img {{
            width: 110px; height: 140px;
            object-fit: cover; border-radius: 4px;
        }}

        /* Sections */
        section {{ margin-bottom: 20px; }}
        h2 {{
            font-size: 13px; font-weight: 700;
            color: #000; text-transform: uppercase;
            letter-spacing: 1px;
            border-bottom: 2px solid #000;
            padding-bottom: 4px; margin-bottom: 12px;
        }}

        /* Entries */
        .entry {{ margin-bottom: 14px; }}
        .entry-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
        .role {{ font-weight: 600; font-size: 13px; }}
        .entry-meta {{ font-size: 11px; color: #333; margin-bottom: 4px; }}
        .entry-notes {{ font-size: 11px; color: #555; margin-top: 2px; }}

        ul {{ padding-left: 20px; margin-top: 4px; }}
        li {{ font-size: 12px; margin-bottom: 3px; }}

        /* Tech stack */
        .stack-row {{ font-size: 12px; margin-bottom: 3px; }}

        /* Language rows */
        .lang-row {{ font-size: 12px; margin-bottom: 3px; }}

        /* Projects */
        .project-link {{
            font-size: 11px; color: #333;
            text-decoration: none; margin-left: 4px;
        }}
        .tech-label {{ font-size: 11px; color: #555; margin-top: 2px; }}

        p {{ font-size: 12px; margin-bottom: 4px; }}

        /* Print */
        @media print {{
            body {{ padding: 20px; max-width: 100%; }}
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="header-info">
            <h1>{p['name']}</h1>
            <div class="title">{p['title']}</div>
            <div class="contact">{contact_html}</div>
        </div>
        <div class="header-portrait">
            <img src="portrait.jpg" alt="Portrait">
        </div>
    </div>

    <section>
        <h2>{headers['summary']}</h2>
        <p>{data['summary']}</p>
    </section>

    <section>
        <h2>{headers['experience']}</h2>
        {exp_html}
    </section>

    <section>
        <h2>{headers['projects']}</h2>
        {proj_html}
    </section>

    <section>
        <h2>{headers['tech_stack']}</h2>
        {ts_html}
    </section>

    <section>
        <h2>{headers['education']}</h2>
        {edu_html}
    </section>

    {awards_html}
    {cert_html}
    {extra_html}
    {lang_html}
</body>
</html>"""

    output_path.write_text(html, encoding="utf-8")
    print(f"  ✓ HTML: {output_path.name}")


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 50)
    print("  CV Generator — EN / ES / KO")
    print("  (Manual translations, no machine translation)")
    print("=" * 50)

    # Load data
    if not DATA_FILE.exists():
        print(f"\n✗ Error: {DATA_FILE} not found.")
        print("  Please create cv_data.json with your CV information.")
        sys.exit(1)

    with open(DATA_FILE, "r", encoding="utf-8") as f:
        raw_data = json.load(f)

    print(f"\n✓ Loaded data for: {raw_data['personal']['name']}")

    # Create output directory
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Copy portrait to output for HTML
    if PORTRAIT_FILE.exists():
        shutil.copy2(str(PORTRAIT_FILE), str(OUTPUT_DIR / "portrait.jpg"))

    # Generate for each language
    for lang_code, lang_name in LANGUAGES.items():
        print(f"\n--- {lang_name} ---")

        # Translate
        cv_data = translate_cv_data(raw_data, lang_code)

        # Generate all 3 formats
        base_name = f"cv_{lang_code}"
        generate_docx(cv_data, lang_code, OUTPUT_DIR / f"{base_name}.docx")
        generate_pdf(cv_data, lang_code, OUTPUT_DIR / f"{base_name}.pdf")
        generate_html(cv_data, lang_code, OUTPUT_DIR / f"{base_name}.html")

    print(f"\n{'=' * 50}")
    print(f"  ✓ Done! 9 files generated in: {OUTPUT_DIR}/")
    print(f"{'=' * 50}")
    print("\nFiles:")
    for f in sorted(OUTPUT_DIR.glob("cv_*")):
        print(f"  {f.name}")


if __name__ == "__main__":
    main()
